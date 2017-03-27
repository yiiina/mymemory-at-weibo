#-*-coding:utf8-*-


from bs4 import BeautifulSoup

import requests


from docx import Document


def parse_text(hlxml,cookie,weibo_order,else_or_self):
    soup=BeautifulSoup(hlxml, "lxml")
    div=soup.find_all('div',class_="c")
    if(else_or_self=='self'):# 自己账户转发微博节点含有5个a节点，别人账户含有3个a节点，不考虑a[转发]节点
        del_num=-6
    else:
        del_num=-4
    text_to_save_doc=''
    document = Document('/Users/Yina/Desktop/mymemory-at-weibo/weibo_to_doc/demo.docx')
    for each in div:
        if len(each.find_all('div'))==1:#原创微博
            text_time=each.select('span[class="ct"]')[0].get_text()
            text_time=text_time.replace("\n", "")
            text_time=text_time.replace(" ", "")
            text_time=text_time.replace('\t','')
            text_time=text_time.strip()
            original_node=each.select('span[class="ctt"]')[0]
            text=original_node.get_text()#
            a_number=len(original_node.select('a[href]'))
            if(a_number!=0):
                #print('我在测试ing',original_node.select('a[href]')[-1].string)
                if("\n      全文\n     " == original_node.select('a[href]')[-1].string):#获取全文
                    #print(original_node.select('a[href]')[-1].contents)
                    all_text_link='http://weibo.cn'+original_node.select('a[href]')[-1]['href']
                    all_text_lxml = requests.get(all_text_link,cookies=cookie).content.decode('utf-8')
                    all_text_soup=BeautifulSoup(all_text_lxml, "lxml")
                    all_text_text=all_text_soup.select('span[class="ctt"]')[0].get_text()
                    if(all_text_text==None):
                        all_text_text=all_text_soup.select('span[class="ctt"]')[0].next_element.get_text()
                    if all_text_text[0]==':':
                        all_text_text=all_text_text[1:]
                    text=all_text_text
            text=text.replace("\n", "")
            text=text.replace(" ", "")
            text=text.replace('\t','')
            text=text.strip()
            text='%d、'%weibo_order+text_time+'\n'+text+'\n'
            weibo_order=weibo_order+1
            document.add_paragraph(text)
            text_to_save_doc=text_to_save_doc+text
            #print(text)

        if len(each.find_all('div'))==2:#转发微博
            text_time=each.select('span[class="ct"]')[0].get_text()
            text_time=text_time.replace("\n", "")
            text_time=text_time.replace(" ", "")
            text_time=text_time.replace('\t','')
            text_time=text_time.strip()
            #text=each.select('span[class="ctt"]')[0].string#仅仅只能获取最简单的内容，不含链接，定位，全文，换行
            transmit_node=each.find_all('div')[1]#转发内容的div节点,即第二个div节点
            #transmit_text=transmit_node.select('span[class="cmt"]')[0].get_text()#转发理由：
            #a=transmit_node.select('span[class="cmt"]')[0].contents
            #del transmit_node('a')
            i=-1
            del_node=each.select('span[class="ct"]')[0]
            del_node.decompose()
            temp_node=transmit_node.select('a')

            while(i!=del_num):
                temp_node[i].decompose()
                i=i-1

            if(len(transmit_node.select('a'))!=0):
                if "赞" in transmit_node.select('a')[-1].string:
                    transmit_node.select('a')[0].decompose()
            if(len(transmit_node.select('span[class="cmt"]'))!=1):
                transmit_node.select('span[class="cmt"]')[0].decompose()

            transmit_text=transmit_node.get_text()#转发理由：***
            transmit_text=transmit_text.replace("\n", "")
            transmit_text=transmit_text.strip()
            transmit_text=transmit_text.replace(" ", "")
            original_weibo_node=each.find_all('div')[0]#被转发内容的div节点,即第一个div节点
            transmit_name_text=original_weibo_node.select('span[class="cmt"]')[0].get_text()#转发了**的微博：
            transmit_name_text=transmit_name_text.replace("\n", "")
            transmit_name_text=transmit_name_text.replace(" ", "")
            #original_text=original_weibo_node.select('span[class="ctt"]')[0].get_text()
            original_node=original_weibo_node.select('span[class="ctt"]')[0]
            original_text=original_node.get_text()
            a_number=len(original_node.select('a[href]'))

            if(a_number!=0):
                #print('我在测试ing',original_node.select('a[href]')[-1].string)
                if("\n      全文\n     " == original_node.select('a[href]')[-1].string):#获取全文
                    all_text_link='http://weibo.cn'+original_node.select('a[href]')[-1]['href']
                    all_text_lxml = requests.get(all_text_link,cookies=cookie).content.decode('utf-8')
                    all_text_soup=BeautifulSoup(all_text_lxml, "lxml")
                    all_text_text=all_text_soup.select('span[class="ctt"]')[0].get_text()
                    if(all_text_text==None):
                        all_text_text=all_text_soup.select('span[class="ctt"]')[0].next_element.get_text()
                    original_text=all_text_text
            original_text=original_text.replace("\n", "")
            original_text=original_text.replace(" ", "")

            if original_text[0]==':':
                original_text=original_text[1:]
            text='%d、'%weibo_order+text_time+'\n'+transmit_text+'\n'+transmit_name_text+original_text
            weibo_order=weibo_order+1

            document.add_paragraph(text)

            text_to_save_doc=text_to_save_doc+text+'\n'

    del div,soup
    document.save('/Users/Yina/Desktop/mymemory-at-weibo/weibo_to_doc/demo.docx')
    return text_to_save_doc,weibo_order





